"""Document upload authorization tests (spec #3)."""

from __future__ import annotations

import io

from sqlalchemy import select

from app.models.document import Document, DocumentChunk
from app.models.user import UserRole
from app.rag.ingestion import CHUNK_MAX_CHARS
from tests.conftest import auth_header


def _ingest_text(client, token, payload):
    return client.post("/api/documents/text", headers=auth_header(token), json=payload)


def _minimal_pdf(text: str) -> bytes:
    """Build a small one-page text PDF without adding a test-only dependency."""
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT\n/F1 12 Tf\n72 720 Td\n({escaped}) Tj\nET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    pdf.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(pdf)


def test_officer_upload_defaults_private(client, make_user, token_for):
    make_user("officer@example.com", role=UserRole.user, department_name="Finance")
    resp = _ingest_text(
        client,
        token_for("officer@example.com"),
        {"title": "Personal note", "text": "private content"},
    )
    assert resp.status_code == 201
    assert resp.json()["visibility"] == "PRIVATE"
    assert resp.json()["access_scope"] == []


def test_department_admin_uploads_for_own_department(client, make_user, token_for):
    make_user("mgr@example.com", role=UserRole.manager, department_name="Finance")
    resp = _ingest_text(
        client,
        token_for("mgr@example.com"),
        {"title": "Finance Note", "text": "The approval limit is five crore."},
    )
    assert resp.status_code == 201, resp.text
    assert resp.json()["visibility"] == "PRIVATE"
    assert resp.json()["access_scope"] == []


def test_department_admin_cannot_upload_for_other_department(client, make_user, token_for):
    """A department admin cannot smuggle another department's scope."""
    make_user("mgr2@example.com", role=UserRole.manager, department_name="Finance")
    resp = _ingest_text(
        client,
        token_for("mgr2@example.com"),
        {"title": "Cross", "text": "x", "access_scope": ["hr"]},
    )
    assert resp.status_code == 403


def test_department_admin_cannot_choose_upload_department(client, make_user, token_for, db):
    """A normal uploader cannot supply a department, even if the id is guessed."""
    make_user("mgr3@example.com", role=UserRole.manager, department_name="Finance")
    # department_id 999 does not exist / is not theirs — must be ignored, not trusted.
    resp = _ingest_text(
        client,
        token_for("mgr3@example.com"),
        {"title": "Owned", "text": "content", "department_id": 999},
    )
    assert resp.status_code == 403


def test_admin_can_upload_for_any_department(client, make_user, token_for, db):
    from sqlalchemy import select

    from app.models.department import Department

    hr_id = db.scalar(select(Department).where(Department.name == "HR")).id
    make_user("boss@example.com", role=UserRole.admin, department_name="Administration")
    resp = client.post(
        "/api/documents/text",
        headers=auth_header(token_for("boss@example.com")),
        json={"title": "HR Doc by Admin", "text": "hr content", "department_id": hr_id},
    )
    assert resp.status_code == 201, resp.text
    assert resp.json()["access_scope"] == ["hr"]
    assert resp.json()["visibility"] == "DEPARTMENT"


def test_document_list_is_scope_filtered(client, make_user, token_for):
    # Admin seeds an HR-only and a common doc.
    admin_token = token_for_admin = None
    make_user("admin2@example.com", role=UserRole.admin, department_name="Administration")
    admin_token = token_for("admin2@example.com")
    client.post(
        "/api/documents/text",
        headers=auth_header(admin_token),
        json={"title": "HR Only", "text": "hr stuff", "access_scope": ["hr"]},
    )
    client.post(
        "/api/documents/text",
        headers=auth_header(admin_token),
        json={"title": "Common Doc", "text": "shared stuff", "access_scope": ["common"]},
    )

    # A finance user sees the common doc but not the HR-only doc.
    make_user("finlist@example.com", role=UserRole.user, department_name="Finance")
    resp = client.get(
        "/api/documents?view=shared",
        headers=auth_header(token_for("finlist@example.com")),
    )
    assert resp.status_code == 200
    titles = {d["title"] for d in resp.json()}
    assert "Common Doc" in titles
    assert "HR Only" not in titles


def test_multipart_upload_is_indexed_and_used_in_answer(client, make_user, token_for, db):
    """Prove the browser upload path reaches retrieval and answer generation."""
    make_user("knowledge.manager@example.com", role=UserRole.manager, department_name="Finance")
    token = token_for("knowledge.manager@example.com")
    upload = client.post(
        "/api/documents",
        headers=auth_header(token),
        data={"title": "Emergency Procurement Rule"},
        files={
            "file": (
                "emergency-procurement.txt",
                b"Emergency Procurement\n\nThe emergency procurement threshold is exactly 47 lakh rupees.",
                "text/plain",
            )
        },
    )
    assert upload.status_code == 201, upload.text
    uploaded = upload.json()
    assert uploaded["status"] == "ingested"
    assert uploaded["document_type"] == "txt"
    assert uploaded["chunk_count"] == 1

    chunks = db.scalars(
        select(DocumentChunk).where(DocumentChunk.document_id == uploaded["id"])
    ).all()
    assert len(chunks) == uploaded["chunk_count"]
    assert chunks[0].embedding

    answer = client.post(
        "/api/rag/query",
        headers=auth_header(token),
        json={"question": "What is the emergency procurement threshold?"},
    )
    assert answer.status_code == 200, answer.text
    body = answer.json()
    assert uploaded["id"] in body["documents_used"]
    assert "47 lakh" in body["answer"].lower()
    assert any(citation["document_id"] == uploaded["id"] for citation in body["citations"])


def test_text_ingest_auto_converts_long_input_to_txt_file(client, make_user, token_for):
    make_user("longtext.manager@example.com", role=UserRole.manager, department_name="Finance")
    token = token_for("longtext.manager@example.com")
    long_text = "Long form policy sentence " * 12

    response = _ingest_text(
        client,
        token,
        {"title": "Long Form Note", "text": long_text},
    )

    assert response.status_code == 201, response.text
    body = response.json()
    assert body["document_type"] == "txt"
    assert body["source_filename"] == "Long Form Note.txt"
    assert body["chunk_count"] >= 1


def test_upload_rejects_unsupported_and_empty_files(client, make_user, token_for, db):
    make_user("upload.manager@example.com", role=UserRole.manager, department_name="Finance")
    headers = auth_header(token_for("upload.manager@example.com"))

    unsupported = client.post(
        "/api/documents",
        headers=headers,
        data={"title": "Binary"},
        files={"file": ("payload.exe", b"not a document", "application/octet-stream")},
    )
    assert unsupported.status_code == 415

    empty = client.post(
        "/api/documents",
        headers=headers,
        data={"title": "Empty"},
        files={"file": ("empty.txt", b"", "text/plain")},
    )
    assert empty.status_code == 422
    assert db.scalars(select(Document)).all() == []


def test_oversized_paragraph_is_split_into_bounded_chunks(client, make_user, token_for, db):
    make_user("chunk.manager@example.com", role=UserRole.manager, department_name="Finance")
    text = ("Procurement policy sentence with searchable terms. " * 150).encode()
    upload = client.post(
        "/api/documents",
        headers=auth_header(token_for("chunk.manager@example.com")),
        data={"title": "Long Policy"},
        files={"file": ("long-policy.txt", text, "text/plain")},
    )
    assert upload.status_code == 201, upload.text
    uploaded = upload.json()
    assert uploaded["chunk_count"] > 1

    chunks = db.scalars(
        select(DocumentChunk).where(DocumentChunk.document_id == uploaded["id"])
    ).all()
    assert chunks
    assert max(len(chunk.text) for chunk in chunks) <= CHUNK_MAX_CHARS


def test_supported_structured_files_are_parsed(client, make_user, token_for):
    from docx import Document as DocxDocument
    from openpyxl import Workbook

    make_user("formats.manager@example.com", role=UserRole.manager, department_name="Finance")
    headers = auth_header(token_for("formats.manager@example.com"))

    docx_buffer = io.BytesIO()
    docx = DocxDocument()
    docx.add_heading("Travel Policy", level=1)
    docx.add_paragraph("Travel reimbursement requires manager approval.")
    docx.save(docx_buffer)

    xlsx_buffer = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Limits"
    sheet.append(["Category", "Limit"])
    sheet.append(["Travel", "25000"])
    workbook.save(xlsx_buffer)

    uploads = [
        ("travel.pdf", _minimal_pdf("Travel reimbursement requires manager approval.")),
        ("travel.docx", docx_buffer.getvalue()),
        ("limits.xlsx", xlsx_buffer.getvalue()),
        ("limits.csv", b"Category,Limit\nTravel,25000\n"),
    ]
    for filename, data in uploads:
        response = client.post(
            "/api/documents",
            headers=headers,
            data={"title": filename},
            files={"file": (filename, data, "application/octet-stream")},
        )
        assert response.status_code == 201, response.text
        assert response.json()["chunk_count"] >= 1
