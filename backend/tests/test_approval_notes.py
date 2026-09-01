"""Approval Note + ONLYOFFICE integration tests (the required 18 cases)."""

from __future__ import annotations

import hashlib
import io

import jwt
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.config import settings
from app.models.approval_note import CompanyDocumentTemplate, TemplateType
from app.models.user import UserRole
from app.services import approval_note_service, docx_populate, onlyoffice_service, template_service
from tests.conftest import auth_header

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
OO_SECRET = "test-onlyoffice-secret"


# --------------------------------------------------------------------------- #
# Helpers / fixtures
# --------------------------------------------------------------------------- #
def _letterhead_docx(with_placeholders: bool = True) -> bytes:
    doc = Document()
    doc.add_heading("ACME CORP — OFFICIAL LETTERHEAD", level=1)
    if with_placeholders:
        doc.add_paragraph("{{APPROVAL_NOTE_TITLE}}")
        doc.add_paragraph("{{APPROVAL_NOTE_CONTENT}}")
        doc.add_paragraph("Prepared by: {{PREPARED_BY}} on {{DATE}}")
    else:
        doc.add_paragraph("Registered office: 1 Example Road.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for section in doc.sections:
        parts += [p.text for p in section.header.paragraphs]
        parts += [p.text for p in section.footer.paragraphs]
    return "\n".join(parts)


def _xml_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(node.text or "" for node in doc.element.body.iter(qn("w:t")))


@pytest.fixture(autouse=True)
def _oo_secret(monkeypatch):
    monkeypatch.setattr(settings, "onlyoffice_jwt_secret", OO_SECRET)
    monkeypatch.setattr(settings, "onlyoffice_url", "http://localhost:8085")
    # Deterministic content so tests never call the real LLM.
    monkeypatch.setattr(
        approval_note_service,
        "generate_content",
        lambda *a, **k: "This is the generated approval note body.\nRecommendation: approve.",
    )


@pytest.fixture
def admin(make_user):
    return make_user("admin@example.com", role=UserRole.admin, department_name="Administration")


@pytest.fixture
def admin_token(admin, token_for):
    return token_for("admin@example.com")


def _upload_letterhead(client, token, data=None, filename="letterhead.docx"):
    return client.post(
        "/api/admin/approval-notes/letterhead",
        headers=auth_header(token),
        files={"file": (filename, data if data is not None else _letterhead_docx(), DOCX_MIME)},
    )


# --------------------------------------------------------------------------- #
# 1-4: letterhead upload + isolation
# --------------------------------------------------------------------------- #
def test_admin_can_upload_letterhead(client, admin_token):
    resp = _upload_letterhead(client, admin_token)
    assert resp.status_code == 201, resp.text
    assert resp.json()["version"] == 1 and resp.json()["is_active"] is True


def test_non_admin_cannot_upload_letterhead(client, make_user, token_for):
    make_user("plain@example.com", role=UserRole.user)
    resp = _upload_letterhead(client, token_for("plain@example.com"))
    assert resp.status_code == 403


def test_invalid_extension_rejected(client, admin_token):
    resp = _upload_letterhead(client, admin_token, data=b"not a docx", filename="evil.txt")
    assert resp.status_code == 415


def test_company_isolation_on_template(client, admin, admin_token, db):
    # A template belonging to a different company must be invisible.
    other = CompanyDocumentTemplate(
        company_id=999,
        template_type=TemplateType.approval_note_letterhead,
        name="Other Co",
        original_filename="other.docx",
        storage_key="ff" * 16,
        is_active=True,
    )
    db.add(other)
    db.commit()
    assert template_service.get_active_letterhead(db, admin) is None  # not company 999's


# --------------------------------------------------------------------------- #
# 5-6: approval note types
# --------------------------------------------------------------------------- #
def test_admin_can_create_custom_type(client, admin_token):
    resp = client.post(
        "/api/admin/approval-notes/types",
        headers=auth_header(admin_token),
        json={"name": "Emergency Procurement Approval Note"},
    )
    assert resp.status_code == 201
    assert resp.json()["name"] == "Emergency Procurement Approval Note"


def test_inactive_type_cannot_be_selected(client, admin, admin_token, make_user, token_for, db):
    _upload_letterhead(client, admin_token)
    t = client.post(
        "/api/admin/approval-notes/types",
        headers=auth_header(admin_token),
        json={"name": "Retired Note"},
    ).json()
    client.patch(
        f"/api/admin/approval-notes/types/{t['id']}",
        headers=auth_header(admin_token),
        json={"is_active": False},
    )
    make_user("u1@example.com", role=UserRole.user, department_name="Engineering")
    resp = client.post(
        "/api/approval-notes",
        headers=auth_header(token_for("u1@example.com")),
        json={"approval_note_type_id": t["id"]},
    )
    assert resp.status_code == 400


# --------------------------------------------------------------------------- #
# 7-11: create workflow, copy, master unchanged, placeholders
# --------------------------------------------------------------------------- #
@pytest.fixture
def created_note(client, admin, admin_token, make_user, token_for, db):
    from app.services.approval_note_type_service import seed_default_types

    _upload_letterhead(client, admin_token)
    seed_default_types(db)  # lifespan seeding doesn't run under TestClient
    type_id = client.get(
        "/api/approval-notes/types", headers=auth_header(admin_token)
    ).json()[0]["id"]
    make_user("author@example.com", role=UserRole.user, department_name="Engineering")
    token = token_for("author@example.com")
    resp = client.post(
        "/api/approval-notes",
        headers=auth_header(token),
        json={"approval_note_type_id": type_id, "title": "CAPEX APPROVAL NOTE"},
    )
    assert resp.status_code == 201, resp.text
    return {"note": resp.json(), "token": token, "type_id": type_id}


def test_create_copies_master_and_master_unchanged(created_note, db, admin):
    template = template_service.get_active_letterhead(db, admin)
    master_bytes = template_service.read_template_bytes(template)
    note = db.get(approval_note_service.ApprovalNote, created_note["note"]["id"])
    note_bytes = approval_note_service.read_note_bytes(note)
    # Copy exists and is a distinct file.
    assert note.storage_key != template.storage_key
    # Master is byte-for-byte unchanged.
    assert hashlib.sha256(master_bytes).hexdigest() == hashlib.sha256(_letterhead_docx()).hexdigest()
    # The working copy differs from the master (placeholders replaced).
    assert note_bytes != master_bytes


def test_placeholders_replaced_and_title_content_inserted(created_note, db):
    note = db.get(approval_note_service.ApprovalNote, created_note["note"]["id"])
    text = _docx_text(approval_note_service.read_note_bytes(note))
    assert "{{APPROVAL_NOTE_TITLE}}" not in text
    assert "{{APPROVAL_NOTE_CONTENT}}" not in text
    assert "CAPEX APPROVAL NOTE" in text  # title inserted
    assert "generated approval note body" in text  # content inserted


def test_fallback_keeps_letterhead_before_generated_content():
    populated = docx_populate.populate_template(
        _letterhead_docx(with_placeholders=False),
        title="CAPEX APPROVAL NOTE",
        content="Generated body",
    )
    text = _docx_text(populated)
    assert text.index("ACME CORP") < text.index("Registered office")
    assert text.index("Registered office") < text.index("CAPEX APPROVAL NOTE")
    assert text.index("CAPEX APPROVAL NOTE") < text.index("Generated body")


def test_placeholders_inside_content_controls_are_populated():
    doc = Document()
    for placeholder in ("{{APPROVAL_NOTE_TITLE}}", "{{APPROVAL_NOTE_CONTENT}}"):
        outer = doc.add_paragraph()
        control = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        inner = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = placeholder
        run.append(text)
        inner.append(run)
        content.append(inner)
        control.append(content)
        outer._p.append(control)
    buffer = io.BytesIO()
    doc.save(buffer)

    populated = docx_populate.populate_template(
        buffer.getvalue(), title="CONTROL TITLE", content="CONTROL BODY"
    )
    text = _xml_text(populated)
    assert "CONTROL TITLE" in text
    assert "CONTROL BODY" in text
    assert "{{APPROVAL_NOTE" not in text


def test_create_without_letterhead_conflicts(client, make_user, token_for, db, admin, admin_token):
    # Seed a type but NO letterhead.
    t = client.post(
        "/api/admin/approval-notes/types",
        headers=auth_header(admin_token),
        json={"name": "No LH Note"},
    ).json()
    make_user("u2@example.com", role=UserRole.user)
    resp = client.post(
        "/api/approval-notes",
        headers=auth_header(token_for("u2@example.com")),
        json={"approval_note_type_id": t["id"]},
    )
    assert resp.status_code == 409


# --------------------------------------------------------------------------- #
# 12: ONLYOFFICE config signed server-side
# --------------------------------------------------------------------------- #
def test_editor_config_is_signed_server_side(created_note, client):
    resp = client.get(
        f"/api/approval-notes/{created_note['note']['id']}/editor-config",
        headers=auth_header(created_note["token"]),
    )
    assert resp.status_code == 200, resp.text
    config = resp.json()["config"]
    assert config["documentType"] == "word"
    assert config["document"]["fileType"] == "docx"
    # The config carries a server-signed JWT that verifies with the secret.
    decoded = jwt.decode(config["token"], OO_SECRET, algorithms=["HS256"])
    assert decoded["document"]["key"] == config["document"]["key"]
    assert config["height"] == "100%" and config["width"] == "100%"
    assert config["editorConfig"]["customization"] == {
        "autosave": True,
        "forcesave": True,
    }


def test_force_save_command_is_signed(created_note, db, monkeypatch):
    import httpx

    note = db.get(approval_note_service.ApprovalNote, created_note["note"]["id"])
    captured = {}

    class FakeResponse:
        def raise_for_status(self):
            return None

        def json(self):
            return {"error": 0, "key": note.onlyoffice_key}

    class FakeClient:
        def __init__(self, **kwargs):
            captured["client"] = kwargs

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return None

        def post(self, url, json):
            captured["url"] = url
            captured["payload"] = json
            return FakeResponse()

    monkeypatch.setattr(httpx, "Client", FakeClient)
    result = onlyoffice_service.force_save(note, note.onlyoffice_key)
    command = jwt.decode(captured["payload"]["token"], OO_SECRET, algorithms=["HS256"])
    assert captured["url"] == "http://localhost:8085/command"
    assert command["c"] == "forcesave" and command["key"] == note.onlyoffice_key
    assert result == {"accepted": True, "error_code": 0, "message": "Save requested."}


# --------------------------------------------------------------------------- #
# 13-15: callback auth, update, cross-document protection
# --------------------------------------------------------------------------- #
def _callback_body(note_id, url, user_id, *, secret=OO_SECRET):
    body = {"status": 2, "url": url, "users": [str(user_id)]}
    body["token"] = jwt.encode(body, secret, algorithm="HS256")
    return body


def test_callback_rejects_invalid_jwt(created_note, client):
    note_id = created_note["note"]["id"]
    cb_token = onlyoffice_service.mint_access_token(note_id, "callback")
    body = _callback_body(note_id, "http://ds/file.docx", 1, secret="WRONG-SECRET")
    resp = client.post(f"/api/onlyoffice/callback/{note_id}?token={cb_token}", json=body)
    assert resp.status_code == 401


def test_callback_saves_new_version(created_note, client, db, monkeypatch):
    note_id = created_note["note"]["id"]
    new_docx = _letterhead_docx()  # any valid docx bytes
    monkeypatch.setattr(onlyoffice_service, "download_saved_file", lambda url: new_docx)
    cb_token = onlyoffice_service.mint_access_token(note_id, "callback")
    body = _callback_body(note_id, "http://ds/cache/file.docx", 1)
    resp = client.post(f"/api/onlyoffice/callback/{note_id}?token={cb_token}", json=body)
    assert resp.status_code == 200 and resp.json() == {"error": 0}
    db.expire_all()
    note = db.get(approval_note_service.ApprovalNote, note_id)
    assert note.document_version == 2
    assert approval_note_service.read_note_bytes(note) == new_docx


def test_callback_token_bound_to_document(created_note, client, db, make_user, token_for, admin_token):
    # Mint a callback token for note A, try to post to a different note id.
    note_a = created_note["note"]["id"]
    cb_token = onlyoffice_service.mint_access_token(note_a, "callback")
    body = _callback_body(9999, "http://ds/file.docx", 1)
    resp = client.post(f"/api/onlyoffice/callback/9999?token={cb_token}", json=body)
    assert resp.status_code == 403  # token note_id != path note_id


# --------------------------------------------------------------------------- #
# 16-17: download + access control
# --------------------------------------------------------------------------- #
def test_download_requires_authorization(created_note, client):
    note_id = created_note["note"]["id"]
    assert client.get(f"/api/approval-notes/{note_id}/download").status_code == 401


def test_user_cannot_access_another_users_note(created_note, client, make_user, token_for):
    note_id = created_note["note"]["id"]
    make_user("other@example.com", role=UserRole.user, department_name="HR")
    resp = client.get(
        f"/api/approval-notes/{note_id}", headers=auth_header(token_for("other@example.com"))
    )
    assert resp.status_code in (403, 404)


# --------------------------------------------------------------------------- #
# 18: existing auth still works
# --------------------------------------------------------------------------- #
def test_existing_auth_still_works(client, make_user, token_for):
    make_user("still@example.com", role=UserRole.user)
    resp = client.get("/api/auth/me", headers=auth_header(token_for("still@example.com")))
    assert resp.status_code == 200 and resp.json()["email"] == "still@example.com"
