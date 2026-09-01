"""Report / Approval-Note automation in the chat flow."""

from __future__ import annotations

import io

from docx import Document

from app.models.user import UserRole
from app.services import approval_note_type_service, report_service, template_service
from tests.conftest import auth_header


def _letterhead() -> bytes:
    doc = Document()
    doc.add_paragraph("{{APPROVAL_NOTE_TITLE}}")
    doc.add_paragraph("{{APPROVAL_NOTE_CONTENT}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_detect_report_intent():
    assert report_service.detect_report_intent("I need a report")
    assert report_service.detect_report_intent("prepare a CAPEX approval note")
    assert report_service.detect_report_intent("please draft a report on procurement")
    assert not report_service.detect_report_intent("take notes on this meeting")
    assert not report_service.detect_report_intent("what is the deployment process?")


def _setup(db, make_user):
    admin = make_user("admin@example.com", role=UserRole.admin, department_name="Administration")
    approval_note_type_service.seed_default_types(db)
    template_service.upload_letterhead(db, admin, filename="lh.docx", data=_letterhead())
    return admin


def test_chat_report_ready_prefills_type_and_params(client, make_user, token_for, db, monkeypatch):
    monkeypatch.setattr(
        report_service, "_extract_details", lambda q, t: ("CAPEX APPROVAL NOTE", {"Amount": "5 crore"})
    )
    _setup(db, make_user)
    resp = client.post(
        "/api/rag/query",
        headers=auth_header(token_for("admin@example.com")),
        json={"query": "prepare a CAPEX approval note for 5 crore"},
    )
    assert resp.status_code == 200, resp.text
    report = resp.json()["report"]
    assert report["status"] == "ready"
    assert report["matched_type_name"].startswith("Capital Expenditure")
    assert report["suggested_parameters"]["Amount"] == "5 crore"
    assert len(report["available_types"]) == 9


def test_chat_report_needs_details_when_type_ambiguous(client, make_user, token_for, db, monkeypatch):
    monkeypatch.setattr(report_service, "_extract_details", lambda q, t: (None, {}))
    _setup(db, make_user)
    resp = client.post(
        "/api/rag/query",
        headers=auth_header(token_for("admin@example.com")),
        json={"query": "I need a report"},
    )
    report = resp.json()["report"]
    assert report["status"] == "needs_details"
    assert report.get("matched_type_id") is None  # excluded when None
    assert len(report["available_types"]) == 9


def test_chat_report_unavailable_without_letterhead(client, make_user, token_for, db, monkeypatch):
    monkeypatch.setattr(report_service, "_extract_details", lambda q, t: (None, {}))
    approval_note_type_service.seed_default_types(db)  # types but no letterhead
    make_user("u@example.com", role=UserRole.user)
    resp = client.post(
        "/api/rag/query",
        headers=auth_header(token_for("u@example.com")),
        json={"query": "I need a report"},
    )
    report = resp.json()["report"]
    assert report["status"] == "unavailable"
    assert report["letterhead_ready"] is False


def test_normal_question_has_no_report(client, make_user, token_for, monkeypatch):
    # Avoid a live LLM call for the non-report path.
    import app.routes.rag as rag_route
    from app.rag.pipeline import RagResult

    monkeypatch.setattr(
        rag_route, "run_agentic_query",
        lambda *a, **k: RagResult(question="q", answer="a", evidence_status="insufficient"),
    )
    make_user("n@example.com")
    resp = client.post(
        "/api/rag/query",
        headers=auth_header(token_for("n@example.com")),
        json={"query": "define hydrostatic pressure"},
    )
    assert resp.status_code == 200
    assert resp.json().get("report") is None
