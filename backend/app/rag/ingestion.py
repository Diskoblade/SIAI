"""Document ingestion pipeline.

    Uploaded document -> parse (Docling / fallback) -> structure-aware chunking
    -> metadata enrichment (inherit access_scope + department) -> embed -> store

Parsing supports PDF, DOCX, XLSX, CSV, Markdown, and plain text. Every chunk
inherits the parent document's authorization metadata.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path

from sqlalchemy.orm import Session

from app.models.document import Document, DocumentChunk, Visibility
from app.rag.embeddings import get_embedder, tokenize
from app.rag.vector_store import get_vector_store

CHUNK_TARGET_CHARS = 900
CHUNK_MAX_CHARS = 1400
CHUNK_OVERLAP_CHARS = 120

SUPPORTED_EXTENSIONS = {".csv", ".docx", ".markdown", ".md", ".pdf", ".txt", ".xlsx"}

_HEADING_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*\s+)?[A-Z][A-Za-z0-9 ,/&()\-]{2,60}$")


@dataclass
class ParsedBlock:
    text: str
    page: int | None = None
    heading: str | None = None


@dataclass
class Chunk:
    text: str
    page: int | None = None
    section: str | None = None
    subsection: str | None = None
    meta: dict = field(default_factory=dict)


class UnsupportedFileType(Exception):
    """Raised when the uploaded file format is not supported."""


class DocumentParseError(Exception):
    """Raised when a supported document cannot be parsed safely."""


class EmptyDocumentError(Exception):
    """Raised when parsing produces no indexable text."""


class VectorIndexingError(Exception):
    """Raised when chunks were persisted but the configured vector store failed."""


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #
def parse_document(filename: str, data: bytes) -> list[ParsedBlock]:
    """Return parsed blocks for a supported file type."""
    if not data:
        raise EmptyDocumentError("The uploaded file is empty.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileType(f"Unsupported file type. Supported extensions: {supported}.")

    if extension in {".txt", ".md", ".markdown"}:
        return _parse_text(_decode_text(data, extension.lstrip(".").upper()))
    if extension == ".csv":
        return _parse_csv(data)

    # Prefer Docling for rich structure-aware parsing when it is installed;
    # otherwise fall back to the lightweight per-format parsers below.
    docling_blocks = _parse_with_docling(filename, data)
    if docling_blocks:
        return docling_blocks

    if extension == ".pdf":
        return _parse_pdf(data)
    if extension == ".docx":
        return _parse_docx(data)
    if extension == ".xlsx":
        return _parse_xlsx(data)
    raise UnsupportedFileType(f"Unsupported file type: {extension}.")


def _parse_with_docling(filename: str, data: bytes) -> list[ParsedBlock] | None:
    """Parse via Docling (Markdown) when available; None to use the fallback."""
    from app.media import docling_parser

    markdown = docling_parser.parse_to_markdown(filename, data)
    if not markdown:
        return None
    blocks = _parse_text(markdown)
    return blocks or None


def _decode_text(data: bytes, label: str) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(f"{label} files must use UTF-8 text encoding.") from exc


def _parse_text(text: str) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        heading = para if _HEADING_RE.match(para) and len(para) < 70 and "\n" not in para else None
        blocks.append(ParsedBlock(text=para, page=None, heading=heading))
    return blocks


def _parse_pdf(data: bytes) -> list[ParsedBlock]:
    try:
        from pypdf import PdfReader  # optional
    except ImportError as exc:  # pragma: no cover - depends on optional dep
        raise UnsupportedFileType(
            "PDF ingestion requires 'pypdf' (pip install pypdf) or Docling."
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise DocumentParseError("Password-protected PDF files are not supported.")
        blocks: list[ParsedBlock] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            for para in re.split(r"\n\s*\n", text):
                para = para.strip()
                if para:
                    blocks.append(ParsedBlock(text=para, page=page_num))
        return blocks
    except DocumentParseError:
        raise
    except Exception as exc:  # pypdf exposes several parser-specific exceptions
        raise DocumentParseError("The PDF could not be read or is damaged.") from exc


def _parse_docx(data: bytes) -> list[ParsedBlock]:
    try:
        import docx  # python-docx, optional
    except ImportError as exc:  # pragma: no cover - depends on optional dep
        raise UnsupportedFileType(
            "DOCX ingestion requires 'python-docx' (pip install python-docx) or Docling."
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
        blocks: list[ParsedBlock] = []
        current_heading: str | None = None
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                current_heading = text
            blocks.append(ParsedBlock(text=text, heading=current_heading))
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                blocks.append(ParsedBlock(text="\n".join(rows), heading=current_heading))
        return blocks
    except Exception as exc:
        raise DocumentParseError("The DOCX file could not be read or is damaged.") from exc


def _clean_cell(value) -> str:
    if value is None:
        return ""
    # Normalize non-breaking spaces and collapse whitespace.
    return re.sub(r"\s+", " ", str(value).replace("\xa0", " ")).strip()


def _cap_cell(value: str, limit: int = 160) -> str:
    """Bound a single cell so one verbose column can't drown the other fields
    (which would push them out of the grader's/answerer's context window)."""
    return value if len(value) <= limit else value[: limit - 1].rstrip() + "…"


def _rows_to_blocks(rows: list[list[str]], *, sheet: str | None) -> list[ParsedBlock]:
    """Turn a grid of cells into one self-describing block per data row.

    Each block reads like ``Course Name: Data Analytics | Credit: 3 | ...`` so a
    row stays meaningful after chunking and retrieves on any of its column
    values. A single-cell row (e.g. a section banner) becomes the heading for the
    rows that follow.
    """
    rows = [r for r in rows if any(cell for cell in r)]
    if not rows:
        return []

    # Header = first row with at least two non-empty cells.
    header: list[str] | None = None
    header_index = 0
    for index, row in enumerate(rows):
        if sum(1 for cell in row if cell) >= 2:
            header = row
            header_index = index
            break

    blocks: list[ParsedBlock] = []
    section = sheet

    if header is None:
        for row in rows:
            text = " | ".join(cell for cell in row if cell)
            if text:
                blocks.append(ParsedBlock(text=text, heading=section))
        return blocks

    headers = [h or f"Column {i + 1}" for i, h in enumerate(header)]
    for row in rows[header_index + 1:]:
        present = [cell for cell in row if cell]
        if not present:
            continue
        if len(present) == 1:  # a section/banner row -> heading for following rows
            section = present[0]
            continue
        pairs = [
            f"{(headers[i] if i < len(headers) else f'Column {i + 1}')}: {_cap_cell(cell)}"
            for i, cell in enumerate(row)
            if cell
        ]
        text = " | ".join(pairs)
        if text:
            prefix = f"[{section}] " if section and section != sheet else ""
            blocks.append(ParsedBlock(text=f"{prefix}{text}", heading=section))
    return blocks


def _parse_xlsx(data: bytes) -> list[ParsedBlock]:
    try:
        from openpyxl import load_workbook  # optional
    except ImportError as exc:  # pragma: no cover - depends on optional dep
        raise UnsupportedFileType(
            "XLSX ingestion requires 'openpyxl' (pip install openpyxl)."
        ) from exc
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        blocks: list[ParsedBlock] = []
        for sheet in wb.worksheets:
            grid = [[_clean_cell(c) for c in row] for row in sheet.iter_rows(values_only=True)]
            blocks.extend(_rows_to_blocks(grid, sheet=sheet.title))
        return blocks
    except UnsupportedFileType:
        raise
    except Exception as exc:
        raise DocumentParseError("The XLSX file could not be read or is damaged.") from exc


def _parse_csv(data: bytes) -> list[ParsedBlock]:
    text = _decode_text(data, "CSV")
    try:
        grid = [[_clean_cell(cell) for cell in row] for row in csv.reader(io.StringIO(text))]
    except csv.Error as exc:
        raise DocumentParseError("The CSV file could not be read.") from exc
    return _rows_to_blocks(grid, sheet=None)


# --------------------------------------------------------------------------- #
# Structure-aware chunking
# --------------------------------------------------------------------------- #
def _split_long_text(text: str) -> list[str]:
    """Split a single oversized block at natural boundaries with light overlap."""
    text = text.strip()
    if len(text) <= CHUNK_MAX_CHARS:
        return [text] if text else []

    parts: list[str] = []
    start = 0
    while start < len(text):
        hard_end = min(start + CHUNK_MAX_CHARS, len(text))
        end = hard_end
        if hard_end < len(text):
            window = text[start:hard_end]
            minimum = min(CHUNK_TARGET_CHARS, len(window) - 1)
            boundaries = [
                window.rfind("\n\n", minimum),
                window.rfind("\n", minimum),
                window.rfind(". ", minimum),
                window.rfind(" ", minimum),
            ]
            boundary = max(boundaries)
            if boundary > 0:
                end = start + boundary + (2 if window[boundary:boundary + 2] == ". " else 0)

        piece = text[start:end].strip()
        if piece:
            parts.append(piece)
        if end >= len(text):
            break

        next_start = max(end - CHUNK_OVERLAP_CHARS, start + 1)
        whitespace = text.find(" ", next_start, end)
        start = whitespace + 1 if whitespace != -1 else next_start
    return parts


def chunk_blocks(blocks: list[ParsedBlock]) -> list[Chunk]:
    """Merge blocks into ~target-sized chunks while preserving section/page."""
    chunks: list[Chunk] = []
    current_section: str | None = None
    buf: list[str] = []
    buf_page: int | None = None

    def flush() -> None:
        nonlocal buf, buf_page
        text = "\n\n".join(buf).strip()
        if text:
            chunks.append(Chunk(text=text, page=buf_page, section=current_section))
        buf = []
        buf_page = None

    expanded_blocks = [
        ParsedBlock(
            text=part,
            page=block.page,
            heading=block.heading if index == 0 else None,
        )
        for block in blocks
        for index, part in enumerate(_split_long_text(block.text))
    ]

    for block in expanded_blocks:
        if block.heading:
            # A heading starts a new section; flush the previous buffer.
            flush()
            current_section = block.heading
        if buf_page is None:
            buf_page = block.page
        candidate = ("\n\n".join(buf + [block.text])).strip()
        if len(candidate) > CHUNK_MAX_CHARS and buf:
            flush()
            buf_page = block.page
        buf.append(block.text)
        if len("\n\n".join(buf)) >= CHUNK_TARGET_CHARS:
            flush()
    flush()
    return chunks


# --------------------------------------------------------------------------- #
# Ingestion entry points
# --------------------------------------------------------------------------- #
def ingest_document(
    db: Session,
    *,
    title: str,
    access_scope: list[str],
    owner_department_id: int | None,
    filename: str,
    data: bytes,
    document_type: str = "document",
    classification: str = "internal",
    created_by: int | None = None,
    owner_user_id: int | None = None,
    visibility: Visibility | None = None,
    memory_category: str | None = None,
) -> Document:
    """Full pipeline: create the Document, parse+chunk+embed, persist, index."""
    blocks = parse_document(filename, data)
    inferred_type = Path(filename).suffix.lower().lstrip(".") or document_type
    return _persist(
        db,
        title=title,
        access_scope=access_scope,
        owner_department_id=owner_department_id,
        filename=filename,
        blocks=blocks,
        document_type=inferred_type,
        classification=classification,
        created_by=created_by,
        owner_user_id=owner_user_id,
        visibility=visibility,
        memory_category=memory_category,
    )


def ingest_text(
    db: Session,
    *,
    title: str,
    text: str,
    access_scope: list[str],
    owner_department_id: int | None,
    document_type: str = "document",
    classification: str = "internal",
    created_by: int | None = None,
    owner_user_id: int | None = None,
    visibility: Visibility | None = None,
    memory_category: str | None = None,
) -> Document:
    """Convenience path for seeding/tests: ingest raw text with no file."""
    return _persist(
        db,
        title=title,
        access_scope=access_scope,
        owner_department_id=owner_department_id,
        filename=f"{title}.txt",
        blocks=_parse_text(text),
        document_type=document_type,
        classification=classification,
        created_by=created_by,
        owner_user_id=owner_user_id,
        visibility=visibility,
        memory_category=memory_category,
    )


def _persist(
    db: Session,
    *,
    title: str,
    access_scope: list[str],
    owner_department_id: int | None,
    filename: str,
    blocks: list[ParsedBlock],
    document_type: str,
    classification: str,
    created_by: int | None,
    owner_user_id: int | None,
    visibility: Visibility | None,
    memory_category: str | None,
) -> Document:
    embedder = get_embedder()
    chunks = chunk_blocks(blocks)
    if not chunks:
        raise EmptyDocumentError(
            "No readable text was found. Scanned PDFs require OCR before upload."
        )
    vectors = embedder.embed_many([chunk.text for chunk in chunks])
    if len(vectors) != len(chunks):
        raise VectorIndexingError("The embedding provider returned an incomplete result.")

    document = Document(
        title=title,
        owner_department_id=owner_department_id,
        owner_user_id=owner_user_id,
        visibility=visibility,
        document_type=document_type,
        classification=classification,
        memory_category=memory_category,
        access_scope=list(access_scope),
        source_filename=filename,
        status="indexing",
        created_by=created_by,
    )
    db.add(document)
    db.flush()  # assign document.id

    for chunk, vector in zip(chunks, vectors, strict=True):
        db.add(
            DocumentChunk(
                document_id=document.id,
                department_id=owner_department_id,
                owner_user_id=owner_user_id,
                visibility=visibility,
                access_scope=list(access_scope),  # inherit authorization metadata
                document_title=title,
                document_type=document_type,
                memory_category=memory_category,
                page=chunk.page,
                section=chunk.section,
                subsection=chunk.subsection,
                text=chunk.text,
                embedding=vector,
                tokens=sorted(set(tokenize(chunk.text))),
            )
        )
    db.commit()
    db.refresh(document)

    # Sync to the vector backend (no-op for the SQLite store; upsert for Qdrant).
    try:
        get_vector_store().upsert(db, document.id)
    except Exception as exc:
        document.status = "index_failed"
        db.commit()
        raise VectorIndexingError("The document was parsed, but vector indexing failed.") from exc

    document.status = "ingested"
    db.commit()
    db.refresh(document)
    return document
