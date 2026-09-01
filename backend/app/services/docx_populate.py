"""Populate an Approval Note DOCX from the letterhead template.

Uses python-docx only (no Word/LibreOffice needed). Placeholders like
``{{APPROVAL_NOTE_TITLE}}`` and ``{{APPROVAL_NOTE_CONTENT}}`` are replaced in the
body, tables, headers and footers. Logos/images/headers/footers/margins are
part of the DOCX package and are preserved because we only rewrite the text of
paragraphs that contain a placeholder.

If the template contains no known placeholders, a graceful fallback inserts the
title near the top of the body and the content beneath it.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Z_]+)\s*\}\}")
TITLE_PLACEHOLDER = "APPROVAL_NOTE_TITLE"
CONTENT_PLACEHOLDER = "APPROVAL_NOTE_CONTENT"


def _all_paragraphs(doc: Document):
    """Yield innermost paragraphs, including content controls and text boxes."""
    roots = [(doc.element.body, doc._body)]
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            roots.append((hf._element, hf))

    seen: set[int] = set()
    paragraph_tag = qn("w:p")
    for root, parent in roots:
        for element in root.iter(paragraph_tag):
            # A drawing/text box can contain another w:p. Process the inner
            # paragraph so inserted content remains inside that shape.
            if sum(1 for _ in element.iter(paragraph_tag)) > 1:
                continue
            marker = id(element)
            if marker in seen:
                continue
            seen.add(marker)
            yield Paragraph(element, parent)


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._p.iter(qn("w:t")))


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace descendant text while preserving the first run's formatting."""
    text_nodes = list(paragraph._p.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(text)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _replace_simple(paragraph: Paragraph, mapping: dict[str, str]) -> bool:
    """Replace non-content placeholders within a single paragraph."""
    full = _paragraph_text(paragraph)
    if "{{" not in full:
        return False

    def _sub(match: re.Match) -> str:
        key = match.group(1)
        if key == CONTENT_PLACEHOLDER:
            return match.group(0)  # handled separately (multi-paragraph)
        return mapping.get(key, match.group(0))

    replaced = _PLACEHOLDER_RE.sub(_sub, full)
    if replaced == full:
        return False
    _set_paragraph_text(paragraph, replaced)
    return TITLE_PLACEHOLDER not in replaced and "{{" in full


def _expand_content(paragraph: Paragraph, content: str) -> None:
    """Replace a {{APPROVAL_NOTE_CONTENT}} paragraph with the content lines."""
    lines = [ln.rstrip() for ln in content.splitlines()] or [""]
    # First line reuses the placeholder paragraph (keeps its style).
    _set_paragraph_text(paragraph, lines[0])
    anchor = paragraph
    for line in lines[1:]:
        anchor = _insert_paragraph_after(anchor, line)


def populate_template(
    template_bytes: bytes,
    *,
    title: str,
    content: str,
    extra_placeholders: dict[str, str] | None = None,
) -> bytes:
    """Return DOCX bytes with the Approval Note title and content applied."""
    doc = Document(io.BytesIO(template_bytes))

    mapping: dict[str, str] = {TITLE_PLACEHOLDER: title}
    if extra_placeholders:
        mapping.update({k: str(v) for k, v in extra_placeholders.items()})

    title_found = False
    content_paragraph: Paragraph | None = None

    for paragraph in _all_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if "{{" not in text:
            continue
        if f"{{{{{TITLE_PLACEHOLDER}}}}}" in text.replace(" ", ""):
            title_found = True
        if CONTENT_PLACEHOLDER in text and content_paragraph is None:
            content_paragraph = paragraph
        _replace_simple(paragraph, mapping)

    if content_paragraph is not None:
        _expand_content(content_paragraph, content)

    # Graceful fallback: neither placeholder present -> insert after the
    # letterhead/front matter instead of pushing the letterhead to page two.
    if not title_found and content_paragraph is None:
        _fallback_insert(doc, title, content)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _fallback_insert(doc: Document, title: str, content: str) -> None:
    body_children = list(doc.element.body.iterchildren())
    last_table = max(
        (index for index, child in enumerate(body_children) if child.tag == qn("w:tbl")),
        default=-1,
    )
    candidate = next(
        (
            Paragraph(child, doc._body)
            for child in body_children[last_table + 1 :]
            if child.tag == qn("w:p")
        ),
        None,
    )

    if candidate is not None and not _paragraph_text(candidate).strip():
        title_para = candidate
        _set_paragraph_text(title_para, title)
    elif doc.paragraphs:
        title_para = _insert_paragraph_after(doc.paragraphs[-1], title)
    else:
        title_para = doc.add_paragraph(title)
    if title_para.runs:
        title_para.runs[0].bold = True
    anchor = title_para
    for line in (content.splitlines() or [""]):
        anchor = _insert_paragraph_after(anchor, line.rstrip())
